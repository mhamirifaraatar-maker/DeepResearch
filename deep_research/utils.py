import sys
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt
from google import genai
import asyncio
from .config import GEMINI_KEY

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("deep_research.log"),
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger(__name__)

def log_error(context: str, error: str):
    """Log an error with context."""
    logger.error(f"[{context}] {error}")

def safe_save(doc: Document, base_path: Path):
    """Save document with a timestamp to avoid overwriting."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = base_path.with_name(f"{base_path.stem}_{timestamp}{base_path.suffix}")
    try:
        doc.save(path)
        logger.info(f"✅ Saved report to {path}")
    except Exception as e:
        log_error("Save", str(e))

def build_doc(report: str) -> Document:
    """Convert markdown report to a Word document."""
    doc = Document()
    
    for line in report.splitlines():
        line = line.strip()
        if not line:
            continue
        
        if line == "---":
            doc.add_page_break()
            continue
        
        if line.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(line[2:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
            continue
        
        if line.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            continue
        
        if line.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(line[4:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            continue
        
        if line.startswith("#### "):
            p = doc.add_heading(level=4)
            run = p.add_run(line[4:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            continue
        
        para = doc.add_paragraph()
        
        # Handle bullet points
        if line.startswith("* "):
            para.style = 'List Bullet'
            line = line[2:]
        
        # Basic markdown parsing for bold and italic
        # Note: This is a simplified parser from the original code
        import re
        tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", line)
        
        for tok in tokens:
            if not tok:
                continue
                
            if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
                run = para.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
                run = para.add_run(tok[1:-1])
                run.italic = True
            else:
                # Fix common LLM markdown error: "*Title*:*" -> "*Title*:"
                # Only replace if it's plain text
                clean_tok = tok.replace(":*", ":")
                run = para.add_run(clean_tok)
            
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    
    return doc

# Initialize Gemini Client lazily
_client = None

def get_client():
    global _client
    if _client is None:
        if not GEMINI_KEY:
            raise ValueError("GEMINI_KEY not found in environment variables.")
        _client = genai.Client(api_key=GEMINI_KEY)
    return _client

async def gemini_complete(prompt: str, max_tokens: int = 6000) -> str:
    """Generate text using Gemini."""
    max_retries = 5
    backoff = 2
    
    for attempt in range(max_retries + 1):
        try:
            client = get_client()
            response = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt,
                config=genai.types.GenerateContentConfig(
                    max_output_tokens=max_tokens,
                    temperature=0.3
                )
            )
            return response.text
        except Exception as e:
            error_str = str(e)
            if "503" in error_str or "429" in error_str:
                if attempt < max_retries:
                    wait_time = backoff * (2 ** attempt)
                    logger.warning(f"Gemini API error ({error_str}). Retrying in {wait_time}s...")
                    await asyncio.sleep(wait_time)
                    continue
            
            log_error("Gemini", error_str)
            return ""
